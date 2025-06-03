import os
import docx
from markdown2 import markdown
from tqdm import tqdm
from ollama import chat
from langchain_community.llms import Ollama

model_1 = Ollama(base_url='http://localhost:11434', model="<<model name>>")  
model_2 = Ollama(base_url='http://localhost:11434', model="<<model name>>")  

messages_base = [
    {
        "role": "system",
        "content": "Sie sind ein IT-Administrator mit über 25 Jahren Erfahrung und arbeiten seit einem Jahr im Bereich der Informationssicherheit."  
    }
]

def evaluate_model(model, prompt_suffix):
    messages = messages_base.copy()  
    messages.append({
        "role": "user",
        "content": f"Erstellen Sie auf Deutsch eine detaillierte Anleitung zur Durchführung folgender Aufgabe: {prompt_suffix}. Verwenden Sie dabei korrekte technische Terminologie im Bereich Cybersicherheit."  
    })

    try:
        response = model.invoke(model=model.model, input=messages)
        return response.strip()
    except Exception as e:
        print(f"An error occurred while connecting to Ollama for {model.model}:", e)
        return ""

def save_output(model_name, doc_obj, format_type="docx"):
    """Saves the model's output to a file. """
    if format_type == "docx":
        doc_obj.save(f"{model_name}_output.docx")
    elif format_type == "html":
        html_text = markdown(doc_obj.paragraphs[0].text)
        with open(f"{model_name}_output.html", 'w') as f:
            f.write(html_text)
    else:
        print(f"Invalid format type: {format_type}")

prompt_1_suffix = "Einrichten eines Zwei-Faktor-Authentifizierungsprozesses (2FA) für Remotezugriff auf Unternehmensressourcen."  
prompt_2_suffix = "Sichern eines lokalen Ordners auf einem Windows-Rechner mit BitLocker Drive Encryption."  

tasks = [
    {"model": model_1, "prompt": prompt_1_suffix, "format": "docx"},
    {"model": model_1, "prompt": prompt_2_suffix, "format": "docx"},
    {"model": model_2, "prompt": prompt_1_suffix, "format": "docx"},
    {"model": model_2, "prompt": prompt_2_suffix, "format": "docx"},
]

# Create document objects for each model
model_docs = {
    model_1.model: docx.Document(),
    model_2.model: docx.Document()
}

def generate_and_save_outputs(tasks):
    for task in tqdm(tasks, desc="Generating outputs"):
        model = task["model"]
        prompt = task["prompt"]
        format_type = task["format"]

        model_name = model.model
        content = evaluate_model(model, prompt)

        # Append content to the respective document
        model_docs[model_name].add_paragraph(content)

    # Save each document after processing all tasks
    for model_name, doc_obj in model_docs.items():
        save_output(model_name, doc_obj, format_type)

    print("Evaluation complete! Results saved.")

generate_and_save_outputs(tasks)
