import docx
from markdown2 import markdown
from tqdm import tqdm
import os
import anthropic
from typing import List

MODEL_NAME = "claude-3-opus-20240229"

def generate_messages(section: str) -> List[dict]:
    """
    Generates a list of user messages for a given section.
    """
    return [
        {
            "role": "user",
            "content": f"{section}"
        }
    ]

client: anthropic.Anthropic = anthropic.Anthropic(api_key=)

print("Generating policy outline...")

system_message = ""

try:
    initial_response = client.messages.create(
        model=MODEL_NAME,
        max_tokens=4000,
        temperature=0.5,
        system=system_message,
        messages=[{"role": "user", "content": "P."}]
    )

    outline = ''.join(block.text for block in initial_response.content if block.type == 'text').strip()
except Exception as e:
    print(f"An error occurred while connecting to the Claude API: {e}")
    exit(1)

print(outline + "\n")
sections = outline.split("\n\n")

word_document = docx.Document()
html_sections = []

for i, section in tqdm(enumerate(sections, start=1), total=len(sections), desc="Processing Section", leave=False):
    messages = generate_messages(section)
    try:
        response = client.messages.create(
            model=MODEL_NAME,
            max_tokens=4000,
            messages=messages  # Removed 'system=system_message' from here
        )
        detailed_info = ''.join(block.text for block in response.content if block.type == 'text').strip()
    except Exception as e:
        print(f"An error occurred: {e}")
        continue
    word_document.add_paragraph(detailed_info)
    word_document.add_paragraph("\n")
    html_sections.append(markdown(detailed_info))

docx_file_path = "e.docx"
with open(docx_file_path, "wb") as doc_file:
    word_document.save(doc_file)

html_file_path = ".html"
with open(html_file_path, "w") as html_file:
    html_file.write("\n".join(html_sections))
